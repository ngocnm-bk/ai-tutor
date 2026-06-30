import docx
from ai_tutor.ingest.processors.docx_processor import extract_text


def test_extract_text_from_docx(tmp_path):
    p = tmp_path / "d.docx"
    doc = docx.Document()
    doc.add_paragraph("Unit 3: Present Simple")
    doc.add_paragraph("Bài tập 1")
    doc.save(str(p))
    out = extract_text(p)
    assert "Present Simple" in out
    assert "Bài tập 1" in out
